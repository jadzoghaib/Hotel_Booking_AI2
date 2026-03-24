"""
Generate the 3-page AI Final Project report as a .docx file.
Run with the project's Python environment.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins (narrow to fit content) ──────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)

# ── Helper: paragraph spacing ─────────────────────────────────────────────────
def set_spacing(para, before=0, after=4, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line:
        pf.line_spacing = Pt(line)

def add_heading(text, level=1, color=RGBColor(0x1F, 0x49, 0x7D)):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = color
        run.font.bold = True
        if level == 1:
            run.font.size = Pt(13)
        elif level == 2:
            run.font.size = Pt(11)
    set_spacing(h, before=8, after=3)
    return h

def add_body(text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    set_spacing(p, before=0, after=3, line=12)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    set_spacing(p, before=0, after=2, line=11)
    return p

# ── Shade table header row ────────────────────────────────────────────────────
def shade_cell(cell, fill_hex="1F497D"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    tcPr.append(shd)

# ══════════════════════════════════════════════════════════════════════════════
# TITLE BLOCK
# ══════════════════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("Hotel Booking Cancellation Prediction & Overbooking Decision Support")
tr.bold      = True
tr.font.size = Pt(15)
tr.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
set_spacing(title, before=0, after=2)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = subtitle.add_run("Final Project — Artificial Intelligence II")
sr.font.size = Pt(10)
sr.italic    = True
set_spacing(subtitle, before=0, after=10)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — PROBLEM DEFINITION
# ══════════════════════════════════════════════════════════════════════════════
add_heading("1. Problem Definition", level=1)

add_body(
    "Hotel revenue management hinges on a fundamental tension: a fully booked hotel "
    "maximises revenue, yet cancellations routinely leave rooms empty on the night. "
    "The global average hotel cancellation rate has risen steadily, reaching roughly "
    "40% in the post-pandemic period, representing a direct loss of potential revenue "
    "and compounding planning difficulties for staffing and procurement. Overbooking — "
    "deliberately accepting more reservations than physical capacity — is the standard "
    "industry counter-measure, but it carries its own cost: if too many guests show up, "
    "the hotel must 'walk' paying guests, incurring compensation costs, reputational "
    "damage, and potential contract penalties."
)

add_heading("Business Decision Supported", level=2)
add_body(
    "This project supports a real-time accept / reject decision for each incoming "
    "booking request. When a new reservation arrives, the system answers: "
    "given current bookings and their individual cancellation likelihoods, does "
    "accepting this guest push the expected number of arrivals above a safe threshold "
    "for their room type? The target variable is therefore binary: "
)
add_bullet("is_canceled = 1  — the booking will be cancelled before arrival")
add_bullet("is_canceled = 0  — the guest will show up")

add_heading("Why Prediction Is Useful", level=2)
add_body(
    "A rule-based overbooking policy (e.g. 'always overbook by 10%') ignores "
    "heterogeneity: a non-refundable direct booking and a free-cancellation online-TA "
    "booking carry vastly different cancellation risks. By assigning each booking an "
    "individual cancellation probability, the hotel can compute expected arrivals with "
    "far greater precision. This per-booking probability feeds directly into a Monte "
    "Carlo simulation that models the full distribution of arrivals, enabling "
    "data-driven, room-type-specific overbooking decisions rather than blunt "
    "hotel-wide rules."
)

add_heading("Operational Workflow", level=2)
add_body(
    "In day-to-day use the pipeline operates as follows: (1) a reservations agent "
    "enters a new booking's attributes into the Streamlit decision tool; "
    "(2) the trained Random Forest model scores the booking and returns a cancellation "
    "probability; (3) that probability is combined with all existing bookings for the "
    "same room type in a 10,000-run Monte Carlo simulation; (4) if expected arrivals "
    "after adding the new guest remain within effective capacity (room-type capacity "
    "plus upgrade buffer), the booking is accepted — otherwise it is rejected. "
    "The entire process runs in under one second, making it suitable for live "
    "front-desk or online-channel use."
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — DATASET & FEATURE ENGINEERING
# ══════════════════════════════════════════════════════════════════════════════
add_heading("2. Dataset, EDA & Feature Engineering", level=1)

add_heading("Dataset Overview", level=2)
add_body(
    "The analysis uses the publicly available Hotel Booking Demand dataset "
    "(Mostipak, Kaggle), containing 119,390 reservations across a City Hotel and "
    "a Resort Hotel spanning 2015–2017. The raw dataset includes 32 columns covering "
    "booking origin, stay characteristics, guest history, room type, pricing, and "
    "the outcome label is_canceled. The overall cancellation rate is 37.0%, "
    "creating a moderate class imbalance (37% positive, 63% negative) that was "
    "preserved through stratified splitting."
)

add_heading("Data Cleaning & Leakage Removal", level=2)
add_body(
    "Three columns were removed as data-leakage risks — they are only populated after "
    "a booking resolves and therefore cannot exist at prediction time: "
    "reservation_status, reservation_status_date, and assigned_room_type. "
    "Four rows with null children values were dropped. "
    "Agent and company NaN values were imputed as 0 (interpretable as 'no agent / "
    "no company involvement'). Complimentary room type P was excluded to maintain "
    "a consistent room-type encoding. The cleaned dataset retains 119,386 rows."
)

add_heading("Exploratory Data Analysis", level=2)
add_body(
    "Key EDA findings: (1) cancellation rates differ materially by deposit type — "
    "non-refundable deposits paradoxically show high cancellation in the raw data, "
    "likely reflecting channel-specific booking patterns; (2) lead time is strongly "
    "positively correlated with cancellation; (3) the City Hotel cancels at a higher "
    "rate than the Resort Hotel; (4) guests from certain origin countries (notably "
    "Portugal) cancel at rates above the global mean. These patterns informed both "
    "feature retention and encoding decisions."
)

add_heading("Feature Engineering & Selection", level=2)
add_body(
    "Categorical encoding followed two strategies: ordinal encoding for room type "
    "(A–L mapped to 1–9, reflecting an assumed quality hierarchy) and month "
    "(mapped to 1–12); one-hot encoding for eight nominal columns "
    "(hotel, meal, country, market_segment, distribution_channel, deposit_type, "
    "customer_type). Rare countries were grouped as 'Other' to reduce dimensionality "
    "and prevent sparse columns from inflating the feature space. "
    "After encoding the dataset expands to 78 candidate features."
)
add_body(
    "Feature selection was driven by a preliminary Random Forest trained solely on "
    "the training partition. Features were ranked by mean impurity reduction "
    "(Gini importance) and a cumulative-importance threshold of 95% was applied: "
    "only features whose cumulative contribution reaches that threshold are retained. "
    "This yielded 40 features — roughly half the candidate set — with the "
    "importance curve visibly flattening beyond that point. The reduction cuts "
    "inference latency and prevents overfitting to noise features without sacrificing "
    "predictive power, an important scalability consideration for a live decision tool."
)
add_body("The most informative features include: ",
         bold_prefix="Top predictors — ")
add_bullet("lead_time — longer lead times signal higher cancellation risk")
add_bullet("deposit_type_Non Refund — counterintuitively associated with cancellations")
add_bullet("adr (average daily rate) — pricing tier correlates with booking intent")
add_bullet("previous_cancellations — guest history is a strong behavioural signal")
add_bullet("total_of_special_requests — more requests correlate with genuine intent to stay")
add_bullet("country of origin, market_segment, distribution_channel")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — MODELS, EVALUATION & STOCHASTIC DECISION SYSTEM
# ══════════════════════════════════════════════════════════════════════════════
add_heading("3. Model Implementation, Evaluation & Decision System", level=1)

add_heading("Model Choice & Justification", level=2)
add_body(
    "Six classifiers were trained on the same 75/25 stratified train/test split "
    "(random_state=42), allowing direct apples-to-apples comparison. Logistic "
    "Regression provides an interpretable linear baseline. Random Forest and "
    "XGBoost were included as the two dominant tree-ensemble paradigms in "
    "structured-data ML. LightGBM adds leaf-wise gradient boosting, which tends "
    "to outperform depth-wise methods on large tabular datasets. A two-layer "
    "MLP (128 → 64 ReLU units) tests whether a neural approach adds value over "
    "ensemble trees. Finally, a soft-voting ensemble combines all four strong "
    "learners to assess whether diversity improves calibration."
)

add_heading("Training Procedure & Validation", level=2)
add_body(
    "Tree-based models were trained on raw (unscaled) features; the MLP and the "
    "logistic regression component within the ensemble were wrapped in a "
    "StandardScaler pipeline to prevent gradient and distance distortion. "
    "Hyperparameter tuning was applied to Random Forest and XGBoost via "
    "RandomizedSearchCV (20 iterations, 3-fold CV, scoring=ROC-AUC). "
    "LightGBM used its high-quality defaults without further tuning. "
    "The MLP used early stopping (monitoring a held-out 10% validation subset) "
    "to prevent over-training. Model generalisation was validated with 5-fold "
    "stratified cross-validation (ROC-AUC) on the full dataset to confirm that "
    "test-set results were not artefacts of a single split."
)

add_heading("Evaluation Metrics", level=2)
add_body(
    "ROC-AUC is the primary ranking metric because the downstream Monte Carlo "
    "system consumes calibrated probabilities, not hard class labels. A model "
    "that ranks cancellation risk correctly across the full probability range is "
    "more valuable than one optimised for a single threshold. Accuracy, Precision, "
    "Recall, and F1 are reported as secondary metrics; the F1-optimal decision "
    "threshold (rather than the default 0.5) is identified on a validation subset "
    "and applied at test time to reflect the asymmetric costs of false positives "
    "(walking a guest) versus false negatives (under-booking)."
)

# ── Performance table ──────────────────────────────────────────────────────────
add_heading("Model Performance Summary", level=2)

headers = ["Model", "Accuracy", "Precision", "Recall", "F1", "Test AUC", "CV AUC (±SD)"]
rows = [
    ["Logistic Regression",  "0.805", "0.803", "0.628", "0.705", "0.885", "0.886 ±0.001"],
    ["Random Forest (tuned)","0.889", "0.879", "0.811", "0.844", "0.958", "0.960 ±0.001"],
    ["XGBoost (tuned)",      "0.879", "0.857", "0.807", "0.831", "0.952", "0.944 ±0.001"],
    ["LightGBM",             "0.877", "0.859", "0.798", "0.827", "0.951", "0.947 ±0.001"],
    ["MLP (Neural Network)", "0.863", "0.846", "0.771", "0.807", "0.939", "0.934 ±0.001"],
    ["Soft-Voting Ensemble", "0.878", "0.876", "0.780", "0.826", "0.951", "—"],
]

tbl = doc.add_table(rows=len(rows)+1, cols=len(headers))
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
for j, h in enumerate(headers):
    cell = tbl.cell(0, j)
    shade_cell(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.bold       = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size  = Pt(9)

# Data rows
for i, row_data in enumerate(rows):
    is_best = (i == 1)  # Random Forest
    for j, val in enumerate(row_data):
        cell = tbl.cell(i+1, j)
        if is_best:
            shade_cell(cell, "D6E4F0")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(val)
        run.font.size = Pt(9)
        if is_best:
            run.bold = True

# Column widths
col_widths = [Cm(3.8), Cm(1.8), Cm(1.8), Cm(1.6), Cm(1.4), Cm(1.8), Cm(2.4)]
for j, w in enumerate(col_widths):
    for row_ in tbl.rows:
        row_.cells[j].width = w

set_spacing(doc.add_paragraph(), before=0, after=4)

add_body(
    "Random Forest achieved the highest test AUC (0.958) and F1 (0.844) and was "
    "selected as the production model, serialised to best_model.pkl via joblib. "
    "The mild train–test AUC gap (0.039) is within acceptable bounds and confirmed "
    "by the tight 5-fold CV score (0.960 ± 0.001), indicating stable generalisation.",
    bold_prefix="Model selection — "
)

add_heading("Stochastic Decision System", level=2)
add_body(
    "The trained model does not make binary accept/reject decisions on its own. "
    "Instead, its output — a per-booking cancellation probability — feeds a "
    "Monte Carlo simulation. For each new booking request, 10,000 independent "
    "scenarios are sampled: every existing booking for the same room type is "
    "treated as a Bernoulli trial (show-up probability = 1 − p_cancel). "
    "Summing the trials in each scenario yields a full empirical distribution of "
    "expected arrivals. The accept rule is: if the expected arrivals after adding "
    "the new guest remain below effective capacity (room-type inventory plus upgrade "
    "buffer — the sum of capacities for all higher-tier room types), the booking is "
    "accepted. This expected-value criterion is complemented by the Monte Carlo "
    "overbooking risk estimate P(arrivals > base capacity), displayed as a histogram "
    "to the reservations manager for full transparency."
)

add_heading("Critical Reflection", level=2)
add_body(
    "Several limitations warrant acknowledgment. First, construct validity: "
    "is_canceled is a necessary proxy for 'guest shows up,' but no-shows (guests "
    "who neither cancel nor arrive) are not captured. Second, distributional shift: "
    "the dataset spans 2015–2017; post-pandemic cancellation behaviour, dynamic "
    "pricing structures, and the growth of flexible-rate OTA bookings may render "
    "historical patterns less predictive. Third, fairness: using country of origin "
    "as a feature may encode demographic proxies and warrants audit before deployment. "
    "Fourth, the 20% cancellation rate assumed in the baseline hotel population is a "
    "configurable parameter — mis-specifying it biases the effective occupancy "
    "estimate and downstream accept/reject decisions. Finally, the upgrade buffer "
    "assumes perfect room-type substitutability within a tier hierarchy, which may "
    "not reflect actual guest preferences or contractual guarantees."
)

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = r"c:\Users\Jad Zoghaib\OneDrive\Desktop\hotel_booking_clean\AI_Final_Report.docx"
doc.save(output_path)
print(f"Report saved to: {output_path}")
